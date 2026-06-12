import pytest

from ash.utils.file_extract import SUPPORTED_EXTENSIONS, to_markdown


def test_md_passthrough(tmp_path):
    f = tmp_path / "spec.md"
    f.write_text("# Title\n\nSome content.")
    assert to_markdown(f) == "# Title\n\nSome content."


def test_txt_passthrough(tmp_path):
    f = tmp_path / "spec.txt"
    f.write_text("Plain text spec.")
    assert to_markdown(f) == "Plain text spec."


def test_unsupported_extension_raises(tmp_path):
    f = tmp_path / "spec.xlsx"
    f.write_bytes(b"fake")
    with pytest.raises(ValueError, match="Unsupported"):
        to_markdown(f)


def test_supported_extensions_set():
    assert {".md", ".txt", ".pdf", ".docx"} == SUPPORTED_EXTENSIONS


def test_pdf_extraction(tmp_path):
    # Write a minimal valid PDF with one page containing "Hello PDF"
    pdf_bytes = _minimal_pdf("Hello PDF")
    f = tmp_path / "spec.pdf"
    f.write_bytes(pdf_bytes)
    result = to_markdown(f)
    assert "Hello PDF" in result


def test_docx_extraction(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("My Heading", level=1)
    doc.add_paragraph("Some paragraph text.")
    f = tmp_path / "spec.docx"
    doc.save(str(f))
    result = to_markdown(f)
    assert "# My Heading" in result
    assert "Some paragraph text." in result


# ── minimal PDF builder ──────────────────────────────────────────────────────

def _minimal_pdf(text: str) -> bytes:
    """Build a tiny but valid PDF containing `text` on one page."""
    stream = f"BT /F1 12 Tf 50 750 Td ({text}) Tj ET"
    stream_bytes = stream.encode()
    stream_len = len(stream_bytes)

    objects = [
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n",
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n",
        (
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R "
            b"/MediaBox [0 0 612 792] /Contents 4 0 R "
            b"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1"
            b" /BaseFont /Helvetica >> >> >> >>\nendobj\n"
        ),
        (
            f"4 0 obj\n<< /Length {stream_len} >>\nstream\n".encode()
            + stream_bytes
            + b"\nendstream\nendobj\n"
        ),
    ]

    body = b"%PDF-1.4\n"
    offsets: list[int] = []
    for obj in objects:
        offsets.append(len(body))
        body += obj

    xref_offset = len(body)
    xref = f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets:
        xref += f"{off:010d} 00000 n \n".encode()

    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n"
    ).encode()

    return body + xref + trailer
